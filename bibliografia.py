import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def criar_documentos_word_por_curso(arquivo_excel, template_word):
    """Gera documentos Word separados para cada curso, inserindo o conteúdo em um template."""
    try:
        # Lê o arquivo Excel
        df = pd.read_excel(arquivo_excel)

        # Garantir que todas as colunas necessárias sejam strings para evitar erro de 'float'
        df['CURSO'] = df['CURSO'].astype(str).fillna('')
        df['EMENTA'] = df['EMENTA'].astype(str).fillna('')
        df['BIBLIOGRAFIA'] = df['BIBLIOGRAFIA'].astype(str).fillna('')
        df['RELATÓRIO ADEQUAÇÃO REFERENDADO'] = df['RELATÓRIO ADEQUAÇÃO REFERENDADO'].astype(str).fillna('')
        df['TIPO'] = df['TIPO'].astype(str).fillna('')
        df['SEMESTRE'] = df['SEMESTRE'].astype(str).fillna('')
        df['DISCIPLINA'] = df['DISCIPLINA'].astype(str).fillna('')

        # Remove espaços extras de todas as linhas da coluna EMENTA
        df['EMENTA'] = df['EMENTA'].apply(lambda x: '\n'.join([linha[3:].lstrip() for linha in x.split('\n')]))

        # Remove linhas sem semestre ou curso
        df = df.dropna(subset=['SEMESTRE', 'CURSO'])

        # Obtém os cursos únicos, garantindo que nenhum seja vazio
        cursos = df['CURSO'].unique()
        print(f"Total de cursos encontrados: {len(cursos)}")  # Debugging

        for curso in cursos:
            if not curso.strip():
                print("Curso ignorado (vazio ou inválido).")
                continue

            print(f"Gerando documento para o curso: {curso}")

            # Filtra disciplinas do curso atual
            df_curso = df[df['CURSO'] == curso].copy()

            # Remove disciplinas duplicadas, considerando nome e semestre
            df_disciplinas = df_curso[['DISCIPLINA', 'SEMESTRE', 'EMENTA', 'RELATÓRIO ADEQUAÇÃO REFERENDADO']].drop_duplicates()

            # Gera o conteúdo do curso
            documento = gerar_conteudo_curso(df_curso, df_disciplinas)

            # Carrega o template Word
            template_documento = Document(template_word)

            # Substitui o placeholder
            substituir_placeholder(template_documento, "<INSERIRCONTEUDO>", documento)

            # Cria um nome de arquivo válido
            arquivo_saida = f"{curso}.docx".replace("/", "_").replace("\\", "_")
            template_documento.save(arquivo_saida)
            print(f"Documento gerado: {arquivo_saida}")

    except FileNotFoundError:
        print(f"Erro: Arquivo '{arquivo_excel}' ou '{template_word}' não encontrado.")
    except Exception as e:
        print(f"Ocorreu um erro inesperado: {e}")

def gerar_conteudo_curso(df_curso, df_disciplinas):
    """Gera o conteúdo do curso como um documento Word formatado."""
    documento = Document()

    # Conversão para numérico e ordenação correta
    semestres = sorted([int(s) for s in df_disciplinas['SEMESTRE'].unique() if s.isdigit() and s != '0'])

    for semestre in semestres:
        p = documento.add_paragraph()
        run = p.add_run(f"{semestre}° Semestre")
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Apenas o semestre centralizado

        # Filtra disciplinas do semestre
        disciplinas = df_disciplinas[df_disciplinas['SEMESTRE'].astype(int) == semestre]

        for _, disciplina in disciplinas.iterrows():
            adicionar_texto_negrito(documento, "Nome da disciplina:")
            documento.add_paragraph(disciplina['DISCIPLINA'])

            adicionar_texto_negrito(documento, "Ementa:")
            documento.add_paragraph(disciplina['EMENTA'])

            # Bibliografias
            bibliografias = df_curso[df_curso['DISCIPLINA'] == disciplina['DISCIPLINA']]
            adicionar_bibliografia(documento, bibliografias, "Básica")
            adicionar_bibliografia(documento, bibliografias, "Complementar")

            adicionar_texto_negrito(documento, "Adequação Referendado:")
            documento.add_paragraph(disciplina['RELATÓRIO ADEQUAÇÃO REFERENDADO'])

    # Adiciona as disciplinas optativas (semestre '0') no final
    optativas = df_disciplinas[df_disciplinas['SEMESTRE'] == '0']
    if not optativas.empty:
        p = documento.add_paragraph()
        run = p.add_run("Disciplinas optativas")
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for _, disciplina in optativas.iterrows():
            adicionar_texto_negrito(documento, "Nome da disciplina:")
            documento.add_paragraph(disciplina['DISCIPLINA'])

            adicionar_texto_negrito(documento, "Ementa:")
            documento.add_paragraph(disciplina['EMENTA'])

            # Bibliografias
            bibliografias = df_curso[df_curso['DISCIPLINA'] == disciplina['DISCIPLINA']]
            adicionar_bibliografia(documento, bibliografias, "Básica")
            adicionar_bibliografia(documento, bibliografias, "Complementar")

            adicionar_texto_negrito(documento, "Adequação Referendado:")
            documento.add_paragraph(disciplina['RELATÓRIO ADEQUAÇÃO REFERENDADO'])

    return documento

def adicionar_texto_negrito(documento, texto):
    """Adiciona um texto em negrito como parágrafo."""
    p = documento.add_paragraph()
    run = p.add_run(texto)
    run.bold = True

def adicionar_bibliografia(documento, df, tipo):
    """Adiciona bibliografia ao documento sem espaço extra entre as referências."""
    adicionar_texto_negrito(documento, f"Bibliografia {tipo}:")
    df['BIBLIOGRAFIA'] = df['BIBLIOGRAFIA'].astype(str).fillna('')  # Converte e remove NaN
    bibliografias = df[df['TIPO'].str.contains(tipo, na=False)]['BIBLIOGRAFIA'].dropna().drop_duplicates()
    for bibliografia in bibliografias:
        p = documento.add_paragraph()
        p.add_run(bibliografia)

def substituir_placeholder(template_documento, placeholder, documento):
    """Substitui o placeholder no documento Word pelo conteúdo."""
    for paragrafo in template_documento.paragraphs:
        if placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(placeholder, "")
            for par in documento.paragraphs:
                novo_par = template_documento.add_paragraph()
                for run in par.runs:
                    novo_run = novo_par.add_run(run.text)
                    novo_run.bold = run.bold
                    novo_run.italic = run.italic

# Caminho dos arquivos (modifique conforme necessário)
arquivo_excel = "/content/ATUALIZAÇÃO BIBLIOGRAFIAS.xlsx"
template_word = "/content/RELAT. DE ADEQUACAO REFERENDADO.docx"

# Chamada da função
criar_documentos_word_por_curso(arquivo_excel, template_word)

